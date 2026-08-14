from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(r"C:\Users\sbhosale\OneDrive - Frontline\Documents\ChatGPT\AT-AI-Test-Automation\output\cv")
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "Sumit_Bhosale_Professional_CV.docx"

# compact_reference_guide preset with named CV overrides:
# Letter portrait; margins 0.58/0.62 in; Arial; 9.4 pt body; 1.05 spacing;
# navy/teal hierarchy; compact real-list geometry; no running header.
NAVY = RGBColor(18, 45, 74)
TEAL = RGBColor(0, 121, 140)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(84, 98, 116)
LIGHT = "D8E2EC"
PALE = "EEF4F7"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.55)
sec.bottom_margin = Inches(0.55)
sec.left_margin = Inches(0.62)
sec.right_margin = Inches(0.62)
sec.header_distance = Inches(0.3)
sec.footer_distance = Inches(0.3)

def set_font(run, size=None, bold=None, color=INK, italic=None, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color is not None: run.font.color.rgb = color

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(9.4)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.05

for style_name, size, before, after in [
    ("Heading 1", 11.5, 8, 3), ("Heading 2", 10.2, 5, 1.5), ("Heading 3", 9.5, 3, 1)
]:
    st = styles[style_name]
    st.font.name = "Arial"
    st._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    st._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = NAVY
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

bullet = styles["List Bullet"]
bullet.font.name = "Arial"
bullet._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
bullet._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
bullet.font.size = Pt(9.25)
bullet.font.color.rgb = INK
bullet.paragraph_format.left_indent = Inches(0.2)
bullet.paragraph_format.first_line_indent = Inches(-0.14)
bullet.paragraph_format.space_after = Pt(1.4)
bullet.paragraph_format.line_spacing = 1.03

def set_bottom_border(paragraph, color="D8E2EC", size="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=50, start=90, bottom=50, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def add_section(title):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = False
    r = p.add_run(title.upper())
    set_font(r, 11.5, True, NAVY)
    set_bottom_border(p, color="9BC2CC", size="10")
    return p

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_font(r, 9.25, color=INK)
    return p

def add_role(company, role, project, dates, meta):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(role)
    set_font(r, 10.4, True, NAVY)
    r = p.add_run(f"  |  {company}")
    set_font(r, 10.1, True, TEAL)
    r = p.add_run(f"\n{project}")
    set_font(r, 9.2, True, INK)
    r = p.add_run(f"  |  {dates}")
    set_font(r, 9.1, False, MUTED)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.keep_with_next = True
    r = p2.add_run(meta)
    set_font(r, 8.9, False, MUTED, italic=True)

# Resume header - customer_pack-inspired named override (ATS-safe, no tables/photos).
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r = p.add_run("SUMIT BHOSALE")
set_font(r, 23, True, NAVY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("SENIOR QUALITY ANALYST")
set_font(r, 11.5, True, TEAL)
r = p.add_run("  |  TEST AUTOMATION  |  API & PERFORMANCE TESTING")
set_font(r, 10.2, True, NAVY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
r = p.add_run("Pune, India  |  +91 77419 00966  |  ")
set_font(r, 9.2, color=MUTED)
r = p.add_run("sumitbhosale0212@gmail.com")
set_font(r, 9.2, True, TEAL)
set_bottom_border(p, color="9BC2CC", size="12")

add_section("Professional Summary")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r = p.add_run(
    "Quality engineering professional with 6+ years of experience delivering web, API, and performance testing across BFSI, investment banking, and healthcare. Hands-on expertise in Selenium WebDriver with Java, Cucumber BDD, TestNG, Rest Assured, Postman, Cypress, and JMeter. Experienced in building maintainable automation frameworks, validating complex trade and healthcare workflows, integrating test suites with CI/CD, and partnering with Agile teams to improve release confidence."
)
set_font(r, 9.45, color=INK)

add_section("Core Expertise")
skills = [
    ("Automation", "Selenium WebDriver, Java, TestNG, Cucumber BDD, Cypress, Page Object Model, data-driven and hybrid frameworks"),
    ("API & Performance", "Rest Assured, Postman, JSON, POJO mapping, serialization/deserialization, JMeter, BlazeMeter"),
    ("CI/CD & Tooling", "Git, GitHub, GitLab CI/CD, Jenkins, Maven, JIRA, Zephyr, Allure, Kibana, Sumo Logic, Sauce Labs"),
    ("Data & Platforms", "SQL, PostgreSQL, MySQL, Windows, Linux, Apple iOS; IntelliJ IDEA, Eclipse, Visual Studio Code"),
    ("Quality Practices", "Functional, regression, sanity, retesting, API, database and performance testing; test planning, execution, traceability, defect reporting, Agile/Scrum")
]
for label, value in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.4)
    r = p.add_run(label + ": ")
    set_font(r, 9.15, True, NAVY)
    r = p.add_run(value)
    set_font(r, 9.15, color=INK)

add_section("Professional Experience")
add_role("Pyramid IT Consulting", "Senior Quality Analyst", "CHESS Replacement - Australian Securities Exchange (ASX)", "Dec 2024 - Jan 2026", "Domain: Banking, Financial Services and Insurance (BFSI)")
for b in [
    "Designed Cucumber BDD feature files and reusable step definitions to automate multiple end-to-end trade flows.",
    "Automated UI validation across TCS BaNCS screens using Selenium WebDriver, Java, and Cucumber BDD.",
    "Automated trade submission and verified acknowledgement receipts through the Quartz Gateway.",
    "Validated trade-capture reports and AR FIX messages for trade submission and cancellation scenarios.",
    "Built parameterized GitLab CI/CD pipelines using environment, trade date, and test-case inputs for flexible execution.",
    "Reviewed Cucumber HTML reports, investigated failures, and communicated actionable results to stakeholders."
]: add_bullet(b)

add_role("Endava Solution", "Quality Analyst", "Health Connect 360 - Cigna, United States", "Dec 2021 - Dec 2024", "Domain: Healthcare | Pharmacy, medical, lab and member-engagement data workflows")
for b in [
    "Developed web and API automation using TestNG, Cucumber BDD, Rest Assured, and data-driven Postman collections.",
    "Created POJO models, test classes, and shared request/response specifications for maintainable API coverage.",
    "Used Allure reporting to capture request, response, and execution logs for faster failure analysis.",
    "Executed performance tests with JMeter, recorded flows with BlazeMeter, and maintained JMX test assets.",
    "Analyzed test failures using Kibana, Sumo Logic, and Sauce Labs; tracked defects and execution through Zephyr Squad.",
    "Prepared test reports and maintained feature files and step definitions aligned with functional specifications."
]: add_bullet(b)

add_role("AdiSoft Technologies", "Quality Analyst", "FNB - FirstRand Limited, South Africa", "Oct 2019 - Jun 2021", "Domain: Investment Banking")
for b in [
    "Reviewed requirements and technical designs, then created detailed test plans, test cases, and traceability coverage.",
    "Built Selenium WebDriver automation using Java, TestNG, Page Object Model, and Maven.",
    "Configured Jenkins jobs to run automated suites and supported repeatable regression execution.",
    "Performed functional, regression, retesting, system, and monthly JMeter performance testing.",
    "Managed defect tracking and reporting, mapped requirements to tests and defects, and participated in test-case reviews.",
    "Contributed to estimation, prioritization, planning, Scrum ceremonies, and continuous improvement of test strategy."
]: add_bullet(b)

add_section("Education & Professional Development")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1.5)
r = p.add_run("Bachelor of Engineering")
set_font(r, 9.6, True, NAVY)
r = p.add_run("  |  Savitribai Phule Pune University  |  2019")
set_font(r, 9.3, color=INK)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Professional coursework: ")
set_font(r, 9.2, True, NAVY)
r = p.add_run("REST API Testing Automation; Selenium WebDriver with Java - Basic to Advanced and Frameworks")
set_font(r, 9.2, color=INK)

# Footer
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(0)
r = fp.add_run("SUMIT BHOSALE  |  SENIOR QUALITY ANALYST")
set_font(r, 7.5, True, MUTED)

# Prevent lone headings and improve pagination.
for p in doc.paragraphs:
    if p.style.name.startswith("Heading"):
        p.paragraph_format.keep_with_next = True

doc.core_properties.title = "Sumit Bhosale - Senior Quality Analyst CV"
doc.core_properties.subject = "Professional CV - Quality Engineering and Test Automation"
doc.core_properties.author = "Sumit Bhosale"
doc.core_properties.keywords = "Quality Analyst, Test Automation, Selenium, Java, API Testing, Rest Assured, Cucumber, JMeter"
doc.save(DOCX)
print(DOCX)
